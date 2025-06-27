import streamlit as st
import os
from dotenv import load_dotenv
from openai import OpenAI
from datetime import datetime
import logging
import re
import io
import httpx
import urllib3
import requests
from bs4 import BeautifulSoup
import urllib.parse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from functools import lru_cache
import time
import base64

# Disable SSL warnings
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# ReportLab imports removed - PDF generation no longer needed

# --- Initialize Session State (add this near the top) ---
if 'results_list' not in st.session_state:
    st.session_state.results_list = [] # Initialize if not already present

# Initialize authentication state
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False

# Initialize performance metrics
if 'performance_metrics' not in st.session_state:
    st.session_state.performance_metrics = {
        'total_searches': 0,
        'total_processing_time': 0,
        'last_search_time': 0,
        'search_times': [],
        'api_response_times': [],
        'document_generation_time': 0,
        'search_history': []
    }

# Initialize search statistics
if 'search_stats' not in st.session_state:
    st.session_state.search_stats = {
        'total_subjects': 0,
        'successful_searches': 0,
        'failed_searches': 0,
        'ofac_hits': 0,
        'risk_alerts': 0
    }

# --- Page Config (MUST be the first Streamlit command) ---
st.set_page_config(
    page_title="Axos Internal AML Demo", 
    layout="centered",
    initial_sidebar_state="collapsed"
)

# Version number
APP_VERSION = "1.72"

# Custom CSS for enhanced UI/UX
st.markdown("""
<style>
    /* Main container styling */
    .main { 
        padding-top: 2rem;
        background-color: #f8f9fa;
        max-width: 70%;
        margin: 0 auto;
    }
    
    /* Container width control */
    .block-container {
        max-width: 70%;
        padding-left: 2rem;
        padding-right: 2rem;
    }
    
    /* Enhanced button styling */
    .stButton>button {
        background-color: #0066CC;
        color: white;
        font-weight: 600;
        border-radius: 8px;
        padding: 0.75rem 2rem;
        border: none;
        transition: all 0.3s ease;
        box-shadow: 0 2px 4px rgba(0,102,204,0.2);
    }
    
    .stButton>button:hover {
        background-color: #0052A3;
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0,102,204,0.3);
    }
    
    /* Primary button special styling */
    .stButton>button[kind="primary"] {
        background: linear-gradient(135deg, #0066CC 0%, #0052A3 100%);
        font-size: 1.1rem;
        padding: 1rem 2.5rem;
    }
    
    /* Metric cards */
    .metric-card {
        background: white;
        padding: 1.5rem;
        border-radius: 12px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        margin-bottom: 1rem;
        border: 1px solid #e0e0e0;
        transition: all 0.3s ease;
    }
    
    .metric-card:hover {
        box-shadow: 0 4px 16px rgba(0,0,0,0.12);
        transform: translateY(-2px);
    }
    
    /* Status cards */
    .status-card {
        background: white;
        padding: 1.25rem;
        border-radius: 10px;
        border-left: 4px solid #0066CC;
        box-shadow: 0 2px 6px rgba(0,0,0,0.06);
        margin-bottom: 0.75rem;
    }
    
    /* Success/Error styling */
    .success-card {
        border-left-color: #28a745;
        background-color: #f8fff9;
    }
    
    .error-card {
        border-left-color: #dc3545;
        background-color: #fff8f8;
    }
    
    /* Text area styling */
    .stTextArea textarea {
        border-radius: 8px;
        border: 2px solid #e0e0e0;
        transition: all 0.3s ease;
    }
    
    .stTextArea textarea:focus {
        border-color: #0066CC;
        box-shadow: 0 0 0 3px rgba(0,102,204,0.1);
    }
    
    /* Select box styling */
    .stSelectbox > div > div {
        border-radius: 8px;
        border: 2px solid #e0e0e0;
        transition: all 0.3s ease;
    }
    
    .stSelectbox > div > div:hover {
        border-color: #0066CC;
    }
    
    /* Progress bar enhancement */
    .stProgress > div > div > div > div {
        background-color: #0066CC;
        background-image: linear-gradient(45deg, rgba(255,255,255,.15) 25%, transparent 25%, transparent 50%, rgba(255,255,255,.15) 50%, rgba(255,255,255,.15) 75%, transparent 75%, transparent);
        background-size: 1rem 1rem;
        animation: progress-bar-stripes 1s linear infinite;
    }
    
    @keyframes progress-bar-stripes {
        from { background-position: 1rem 0; }
        to { background-position: 0 0; }
    }
    
    /* Headers styling */
    h1 {
        color: #1a1a1a;
        font-weight: 700;
        letter-spacing: -0.5px;
    }
    
    h2, h3 {
        color: #333;
        font-weight: 600;
    }
    
    /* Info boxes */
    .stInfo {
        background-color: #e8f4fd;
        border-color: #0066CC;
        border-radius: 8px;
    }
    
    /* Success boxes */
    .stSuccess {
        background-color: #d4edda;
        border-color: #28a745;
        border-radius: 8px;
    }
    
    /* Error boxes */
    .stError {
        background-color: #f8d7da;
        border-color: #dc3545;
        border-radius: 8px;
    }
    
    /* Download button special styling */
    .stDownloadButton > button {
        background-color: #28a745;
        color: white;
        font-weight: 600;
        border-radius: 8px;
        padding: 0.75rem 2rem;
        border: none;
        transition: all 0.3s ease;
    }
    
    .stDownloadButton > button:hover {
        background-color: #218838;
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(40,167,69,0.3);
    }
    
    /* Sidebar styling */
    .css-1d391kg {
        background-color: #f1f3f5;
    }
    
    /* Metric value styling */
    [data-testid="metric-container"] {
        background-color: white;
        border: 1px solid #e0e0e0;
        padding: 1rem;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    
    /* Animation for loading states */
    @keyframes pulse {
        0% { opacity: 1; }
        50% { opacity: 0.5; }
        100% { opacity: 1; }
    }
    
    .loading-pulse {
        animation: pulse 2s ease-in-out infinite;
    }
</style>
""", unsafe_allow_html=True)

# Configure logging with enhanced format
logging.basicConfig(
    level=logging.INFO, 
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler()  # Remove file handler to improve performance
    ]
)

# Function to get client IP address
def get_client_ip():
    """Get the client's IP address from Streamlit context"""
    try:
        # Try to get from modern Streamlit context headers
        try:
            import streamlit as st
            headers = st.context.headers
            if headers:
                # Check for forwarded headers first (for proxy/load balancer scenarios)
                forwarded_for = headers.get('x-forwarded-for') or headers.get('x-real-ip')
                if forwarded_for:
                    return forwarded_for.split(',')[0].strip()
        except (AttributeError, ImportError):
            # Fallback for older Streamlit versions (but this will show deprecation warning)
            try:
                from streamlit.web.server.websocket_headers import _get_websocket_headers
                headers = _get_websocket_headers()
                if headers:
                    # Check for forwarded headers first (for proxy/load balancer scenarios)
                    forwarded_for = headers.get('x-forwarded-for', headers.get('x-real-ip'))
                    if forwarded_for:
                        return forwarded_for.split(',')[0].strip()
            except ImportError:
                pass
        
        # Fallback: try to get from Streamlit runtime
        import streamlit.runtime.scriptrunner as sr
        session_info = sr.get_script_run_ctx()
        if session_info and hasattr(session_info, 'session_id'):
            return f"session_{session_info.session_id[:8]}"
            
    except Exception as e:
        logging.debug(f"Could not determine client IP: {e}")
    
    return "unknown_ip"

# Function to log user requests
def log_user_request(request_type, company_name, model=None, additional_info=None):
    """Log user requests with IP address and details"""
    client_ip = get_client_ip()
    log_message = f"USER_REQUEST - IP: {client_ip} - Type: {request_type} - Company: '{company_name}'"
    
    if model:
        log_message += f" - Model: {model}"
    
    if additional_info:
        log_message += f" - Info: {additional_info}"
    
    logging.info(log_message)

# --- Configuration & API Client Setup ---
load_dotenv()  # Load .env file for local development

@st.cache_data
def load_api_key(key_name):
    """Load API key from environment or secrets with caching"""
    # Try environment variable first
    key = os.getenv(key_name)
    if key and key.strip():
        return key.strip(), "Environment Variable"
    
    # Try Streamlit secrets
    try:
        key = st.secrets.get(key_name)
        if key and key.strip():
            return key.strip(), "Streamlit Secrets"
    except Exception:
        pass
    
    return None, "Not Found"

# Load API keys
PERPLEXITY_API_KEY, perplexity_source = load_api_key('PERPLEXITY_API_KEY')
OPENAI_API_KEY, openai_source = load_api_key('OPENAI_API_KEY')
API_KEY_LOADED_SUCCESSFULLY = bool(PERPLEXITY_API_KEY)
SOURCE_MESSAGE = f"Perplexity Key Source: {perplexity_source}"

# --- Add Debugging Output Early --- 
st.sidebar.info(SOURCE_MESSAGE) # Show where the key was (or wasn't) found
if API_KEY_LOADED_SUCCESSFULLY:
    # Mask key for display
    masked_key = f"{PERPLEXITY_API_KEY[:7]}...{PERPLEXITY_API_KEY[-4:]}" if PERPLEXITY_API_KEY and len(PERPLEXITY_API_KEY) > 11 else "Invalid Key Format"
    st.sidebar.success(f"Perplexity API Key: Loaded ({masked_key})")
else:
    st.sidebar.error("Perplexity API Key: NOT loaded.")

# OpenAI API Key status
if OPENAI_API_KEY:
    masked_openai_key = f"{OPENAI_API_KEY[:7]}...{OPENAI_API_KEY[-4:]}" if len(OPENAI_API_KEY) > 11 else "Invalid Key Format"
    st.sidebar.success(f"OpenAI API Key: Loaded ({masked_openai_key}) - Source: {openai_source}")
else:
    st.sidebar.warning(f"OpenAI API Key: NOT loaded - Source: {openai_source}")
# --- End Debugging Output ---

PERPLEXITY_API_BASE_URL = "https://api.perplexity.ai"

# Create a shared HTTP client for all requests
@st.cache_resource
def get_http_client():
    """Get a shared HTTP client with SSL disabled"""
    return httpx.Client(verify=False)

# Initialize API clients with singleton pattern
@st.cache_resource
def get_perplexity_client():
    """Get or create Perplexity API client"""
    if not PERPLEXITY_API_KEY:
        return None, "No API Key"
    
    try:
        http_client = get_http_client()
        client = OpenAI(
            api_key=PERPLEXITY_API_KEY,
            base_url=PERPLEXITY_API_BASE_URL,
            http_client=http_client
        )
        return client, None
    except Exception as e:
        return None, str(e)

# Initialize clients
openai_client, client_init_error_msg = get_perplexity_client()

# Update sidebar status
if API_KEY_LOADED_SUCCESSFULLY:
    if openai_client:
        st.sidebar.success("API Client Status: Initialized.")
    else:
        st.sidebar.error(f"API Client Status: Failed ({client_init_error_msg})")
else:
    st.sidebar.warning("API Client Status: Not initialized (No API Key).")

# Error handling
if not openai_client and ("AI" in st.session_state.get("search_engine", "") or "Comprehensive" in st.session_state.get("search_engine", "")):
    error_msg = "ERROR: Perplexity API client required but not initialized. "
    if not API_KEY_LOADED_SUCCESSFULLY:
        error_msg += f"API key was not loaded ({SOURCE_MESSAGE}). "
    elif client_init_error_msg:
        error_msg += f"Client initialization failed: {client_init_error_msg}. "
    error_msg += "Please check API Key configuration."
    st.error(error_msg)
    st.stop()

# Removed unused helper functions for cleaner code

# --- Constants ---
NEGATIVE_KEYWORDS = '(arrest OR bankruptcy OR BSA OR conviction OR criminal OR fraud OR trafficking OR lawsuit OR "money laundering" OR OFAC OR Ponzi OR terrorist OR violation OR "honorary consul" OR consul OR "Panama Papers" OR theft OR corruption OR bribery)'

# Preferred domain list for enhanced searches
PREFERRED_DOMAINS = [
    "sec.gov",
    "finra.org", 
    "treasury.gov",
    "oig.gov",
    "cftc.gov",
    "fdic.gov",
    "federalreserve.gov",
    "justice.gov",
    "fbi.gov",
    "fincen.gov",
    "reuters.com",
    "bloomberg.com",
    "wsj.com",
    "ft.com",
    "law360.com"
]

# Perplexity Models - Available for testing
PERPLEXITY_MODELS = {
    "sonar-pro": {
        "name": "Sonar Pro",
        "description": "Fast, efficient AI search with real-time web access",
        "use_case": "Standard research and analysis",
        "max_tokens": 8000
    },
    "sonar-deep-research": {
        "name": "Sonar Deep Research", 
        "description": "Exhaustive research across hundreds of sources with expert-level analysis",
        "use_case": "Comprehensive reports and detailed investigations",
        "max_tokens": 15000
    }
}

# Default models for different use cases
DEFAULT_PERPLEXITY_MODEL = "sonar-pro"
DEFAULT_COMPREHENSIVE_MODEL = "sonar-pro"
DEFAULT_OFAC_FALLBACK_MODEL = "sonar-pro"

# --- Logo Loading Function ---
@st.cache_data
def get_axos_logo_base64():
    """Load and encode the Axos Bank logo"""
    try:
        with open("axos_bank.jpg", "rb") as image_file:
            encoded_string = base64.b64encode(image_file.read()).decode()
        return encoded_string
    except FileNotFoundError:
        # Return a placeholder if logo file is not found
        st.warning("Logo file 'axos_bank.jpg' not found. Please add the logo file to the project directory.")
        return ""
    except Exception as e:
        st.error(f"Error loading logo: {str(e)}")
        return ""

# --- Authentication System ---
# Password configuration - hardcoded for demo purposes
APP_PASSWORD = os.getenv('APP_PASSWORD', 'AML2024secure!')

def check_password():
    """Returns True if the user has entered the correct password."""
    
    def password_entered():
        """Checks whether a password entered by the user is correct."""
        try:
            client_ip = get_client_ip()
        except NameError:
            client_ip = "unknown_ip"
        
        if st.session_state["password"] == APP_PASSWORD:
            st.session_state["authenticated"] = True
            logging.info(f"AUTH_SUCCESS - IP: {client_ip} - User successfully authenticated")
            del st.session_state["password"]  # Don't store password
        else:
            st.session_state["authenticated"] = False
            logging.warning(f"AUTH_FAILED - IP: {client_ip} - Failed authentication attempt")

    if not st.session_state.get("authenticated", False):
        # Show login form - logo first, then title
        # Add Axos Bank logo to top of login screen
        try:
            axos_logo_base64 = get_axos_logo_base64()
            if axos_logo_base64:
                st.markdown(f"""
                <div style="display: flex; justify-content: center; margin: 20px 0;">
                    <img src="data:image/jpeg;base64,{axos_logo_base64}" 
                         style="height: 60px; width: auto;">
                </div>
                """, unsafe_allow_html=True)
        except Exception as e:
            logging.warning(f"Could not load Axos logo on login screen: {e}")
        
        st.markdown("# AI Center of Excellence - Agentic Research Platform")
        st.markdown("### Please enter the access password to continue")
        
        st.text_input(
            "Password", 
            type="password", 
            on_change=password_entered, 
            key="password",
            placeholder="Enter password to access the system"
        )
        
        if "password" in st.session_state and st.session_state["password"]:
            if not st.session_state.get("authenticated", False):
                st.error("❌ Incorrect password. Please try again.")
        
        st.info("💡 This system is restricted to authorized personnel only.")
        return False
    else:
        return True

# Check authentication before proceeding
if not check_password():
    st.stop()

# Add logout button to sidebar for authenticated users
with st.sidebar:
    st.markdown("---")
    if st.button("🚪 Logout", type="secondary"):
        st.session_state.authenticated = False
        logging.info(f"AUTH_LOGOUT - IP: {get_client_ip()} - User logged out")
        st.rerun()

# --- Helper Functions ---
def extract_citations(message, response):
    """Extract and standardize citations from API response"""
    citations = []
    raw_citations = []
    
    if hasattr(message, 'citations') and message.citations:
        raw_citations = message.citations
    elif hasattr(response, 'citations') and response.citations:
        raw_citations = response.citations
    
    for cit in raw_citations:
        citation_dict = {'url': '#', 'title': 'Source'}
        if isinstance(cit, dict):
            citation_dict['url'] = cit.get('url', '#')
            citation_dict['title'] = cit.get('title', cit.get('url', 'Source'))
        elif hasattr(cit, 'url'):
            citation_dict['url'] = getattr(cit, 'url', '#')
            citation_dict['title'] = getattr(cit, 'title', getattr(cit, 'url', 'Source'))
        elif isinstance(cit, str):
            citation_dict['url'] = cit
            citation_dict['title'] = cit
        else:
            citation_dict['title'] = str(cit)
        citations.append(citation_dict)
    
    return citations

# --- Core Functions (Adapted from Flask app) ---

@lru_cache(maxsize=100)
def search_with_perplexity(company_name, model=DEFAULT_PERPLEXITY_MODEL):
    # (This function remains largely the same as in app.py)
    # ... (API call logic, prompt, message structure) ...
    start_time = time.time()  # Track API call time
    client_ip = get_client_ip()
    model_config = PERPLEXITY_MODELS.get(model, {})
    logging.info(f"SEARCH_START - IP: {client_ip} - Perplexity search for '{company_name}' using model '{model}' ({model_config.get('name', 'Unknown Model')})")
    if not openai_client:
        logging.error("OpenAI client (for Perplexity) not initialized.")
        return {"status": "failed", "error": "Perplexity API client not initialized.", "answer": None, "citations": []}
    try:
        # Build domain instruction using preferred domains
        domain_instruction = f"\n\nIMPORTANT: Prioritize information from these authoritative domains first: {', '.join(PREFERRED_DOMAINS)}. Search these regulatory, financial, and news sources before other sources as they provide the most reliable AML-relevant information."
        
        # Updated Prompt: Ask for explicit separation with headings
        prompt = (
            f"Provide a comprehensive AML (Anti-Money Laundering) due diligence assessment for '{company_name}'. {domain_instruction}"
            f"\n\nStructure your response as follows:"
            f"\n\n## Subject Summary"
            f"\nProvide a brief summary of the subject '{company_name}', including PRODUCTS and SERVICES offered, business activities, key executives, and geographic presence."
            f"\n\n### Subject Websites"
            f"\nIMPORTANT: You MUST include this section. List all official websites for the subject as clickable URLs (e.g., https://example.com). If no websites are found, state: 'This search did not identify potential websites for the subject.'"
            f"\n\n## AML Risk Assessment"
            f"\nAnalyze any negative news found regarding this subject, focusing on: {NEGATIVE_KEYWORDS}. "
            f"Organize findings into clear categories such as 'Financial Crimes', 'Regulatory Issues', 'Legal Proceedings', etc. "
            f"For each finding, include when it happened, key parties involved, and current status if available. "
            f"If no relevant negative news is found in a category, state that clearly.\n"
            f"\n\n## Summary & Overall Assessment"
            f"\nProvide a clear summary of key findings identified regarding this entity. "
            f"Summarize the overall situation including:"
            f"\n- Key risk factors identified (if any)"
            f"\n- Notable compliance or regulatory issues"
            f"\n- Current status of any ongoing matters"
            f"\n- Overall risk profile based on available information"
            f"\n\nPresent findings objectively without making business recommendations. "
            f"Use double line breaks between sections. Provide citations as numeric references like [1], [2] etc., within the text where applicable."
        )
        # For Sonar models, system prompts are ignored - combine everything in user message
        if "sonar" in model.lower():
            full_prompt = f"As an expert AML analyst performing subject due diligence, provide comprehensive analysis with clear subject summary, risk assessment with organized categories, and objective findings summary. Use numeric citations [1] and maintain clean formatting with proper section headers.\n\n{prompt}"
            messages = [
                {"role": "user", "content": full_prompt}
            ]
        else:
            # For other models, use system + user structure
            messages = [
                {
                    "role": "system",
                    "content": "You are an expert AML analyst performing subject due diligence. Provide comprehensive analysis with clear subject summary, risk assessment with organized categories, and objective findings summary. Use numeric citations [1] and maintain clean formatting with proper section headers.",
                },
                {"role": "user", "content": prompt},
            ]
        
        # Build API parameters based on model type
        max_tokens = PERPLEXITY_MODELS.get(model, {}).get("max_tokens", 2000)
        api_params = {
            "model": model,
            "messages": messages,
            "max_tokens": max_tokens
        }
        
        # Note: reasoning_effort parameter not yet supported in current OpenAI client version
        # Will be added when the client is updated to support this parameter
        if "deep-research" in model.lower():
            logging.info(f"Deep research model selected - enhanced analysis will be performed")
        
        # Only add temperature for non-Sonar models (real-time search models don't use temperature)
        if "sonar" not in model.lower():
            api_params["temperature"] = 0.1
        
        # Note: web_search_options removed as it's not supported by the current API
        
        logging.info(f"Calling Perplexity API with model: {model} (max_tokens: {max_tokens})...")
        response = openai_client.chat.completions.create(**api_params)
        logging.info(f"Perplexity API call completed using model: {model}")
        
        full_answer_content = None
        citations = []
        recommendation = None
        if response.choices and len(response.choices) > 0:
            message = response.choices[0].message
            if message and message.content:
                full_answer_content = message.content
                # No longer extracting recommendations - providing objective findings only
                recommendation = None
            # Extract citations
            citations = extract_citations(message, response)
                     
        if not full_answer_content:
            full_answer_content = "No summary could be generated by Perplexity."
        
        # Track API response time
        api_time = time.time() - start_time
        if 'performance_metrics' in st.session_state:
            st.session_state.performance_metrics['api_response_times'].append(api_time)
        logging.info(f"Perplexity API response time: {api_time:.2f}s")
            
        return {"status": "success", "error": None, "answer": full_answer_content, "citations": citations, "recommendation": recommendation, "api_time": api_time}

    except Exception as e:
        logging.error(f"Error during Perplexity search for {company_name}: {str(e)}", exc_info=True)
        api_time = time.time() - start_time
        return {"status": "failed", "error": str(e), "answer": None, "citations": [], "recommendation": None, "api_time": api_time}

# Create a shared session for OFAC requests
@st.cache_resource
def get_ofac_session():
    """Get a shared session for OFAC requests"""
    session = requests.Session()
    session.verify = False
    return session

@lru_cache(maxsize=100)
def search_with_ofac(query):
    """Search OFAC sanctions database"""
    start_time = time.time()  # Track OFAC search time
    client_ip = get_client_ip()
    logging.info(f"SEARCH_START - IP: {client_ip} - OFAC search for '{query}'")
    try:
        session = get_ofac_session()
        
        # Get the initial page to extract form data
        initial_url = "https://sanctionssearch.ofac.treas.gov/"
        initial_response = session.get(initial_url, timeout=30)
        
        if initial_response.status_code != 200:
            return f"Failed to access OFAC website. Status code: {initial_response.status_code}"
        
        soup = BeautifulSoup(initial_response.content, 'html.parser')
        
        # Extract form data
        form_data = {}
        
        # Get all hidden fields
        hidden_fields = soup.find_all('input', {'type': 'hidden'})
        for field in hidden_fields:
            name = field.get('name')
            value = field.get('value', '')
            if name:
                form_data[name] = value
        
        # Add search parameters
        form_data.update({
            'ctl00$MainContent$txtLastName': query,
            'ctl00$MainContent$ddlType': '',
            'ctl00$MainContent$txtAddress': '',
            'ctl00$MainContent$txtCity': '',
            'ctl00$MainContent$txtID': '',
            'ctl00$MainContent$txtState': '',
            'ctl00$MainContent$lstPrograms': '',
            'ctl00$MainContent$ddlCountry': '',
            'ctl00$MainContent$ddlList': '',
            'ctl00$MainContent$Slider1': '83',
            'ctl00$MainContent$Slider1_Boundcontrol': '83',
            'ctl00$MainContent$btnSearch': 'Search',
            '__EVENTTARGET': 'ctl00$MainContent$btnSearch',
            '__EVENTARGUMENT': ''
        })
        
        # Remove the button value since we're using __EVENTTARGET
        if 'ctl00$MainContent$btnSearch' in form_data:
            del form_data['ctl00$MainContent$btnSearch']
        
        # Submit the search
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Content-Type': 'application/x-www-form-urlencoded',
            'Referer': initial_url,
            'Origin': 'https://sanctionssearch.ofac.treas.gov'
        }
        
        search_response = session.post(initial_url, data=form_data, headers=headers, timeout=30)
        
        if search_response.status_code != 200:
            return f"Search request failed. Status code: {search_response.status_code}"
        
        # Parse results
        result_soup = BeautifulSoup(search_response.content, 'html.parser')
        
        # Look for results table
        results_table = result_soup.find('table', {'id': 'gvSearchResults'})
        
        if not results_table:
            # Check for "no results" indicators
            results_div = result_soup.find('div', {'id': 'ctl00_MainContent_divResults'})
            if results_div:
                results_text = results_div.get_text()
                if "0 Found" in results_text or "No results" in results_text.lower():
                    return "✅ No matches found in OFAC sanctions database - entity appears clean"
            
            return "❌ Could not parse search results from OFAC website"
        
        # Parse the results table
        rows = results_table.find_all('tr')
        if len(rows) <= 1:  # Only header row or no rows
            return "✅ No matches found in OFAC sanctions database - entity appears clean"
        
        # Extract results count from the page
        results_count = 0
        results_label = result_soup.find('span', {'id': 'ctl00_MainContent_lblResults'})
        if results_label:
            results_text = results_label.get_text()
            count_match = re.search(r'(\d+)\s+Found', results_text)
            if count_match:
                results_count = int(count_match.group(1))
        
        if results_count == 0:
            return "✅ **OFAC CLEAR**: No matches found in OFAC sanctions database\n\n**Summary**: Entity appears clean from sanctions perspective.\n\n**Recommendation**: Proceed with standard due diligence protocols."
        
        # Initialize variables at the start
        all_matches = []
        high_confidence_matches = 0
        
        # Process ALL rows - OFAC table has no header row to skip
        for i, row in enumerate(rows, 1):
            cells = row.find_all('td')
            
            # Must have at least 6 cells for a valid data row
            if len(cells) < 6:
                continue
                
            try:
                # Extract all cell data
                name_cell = cells[0]
                name_link = name_cell.find('a')
                name = name_link.get_text().strip() if name_link else name_cell.get_text().strip()
                
                address = cells[1].get_text().strip()
                entity_type = cells[2].get_text().strip()
                programs = cells[3].get_text().strip()
                list_type = cells[4].get_text().strip()
                score = cells[5].get_text().strip()
                
                # Parse score - handle both "100%" and "100" formats
                score_num = 0
                if score:
                    try:
                        score_clean = score.replace('%', '').strip()
                        score_num = float(score_clean)
                    except:
                        continue
                else:
                    continue
                
                # Apply 83% threshold
                if score_num < 83:
                    continue

                # Track high confidence matches (any match above 83% threshold is significant)
                if score_num >= 83:
                    high_confidence_matches += 1
                
                # Store match info
                score_display = f"{score_num}%" if not score.endswith('%') else score
                all_matches.append({
                    'name': name,
                    'address': address,
                    'entity_type': entity_type,
                    'programs': programs,
                    'list_type': list_type,
                    'score': score_display,
                    'score_num': score_num
                })
                
            except Exception as e:
                continue
        
        # Check if no matches found
        if len(all_matches) == 0:
            return "✅ **OFAC CLEAR**: No matches found in OFAC sanctions database\n\n**Summary**: Entity appears clean from sanctions perspective.\n\n**Recommendation**: Proceed with standard due diligence protocols."
        
        # Sort matches by score (highest first)
        all_matches.sort(key=lambda x: x['score_num'], reverse=True)
        
        # Format the results
        result_text = f"🚨 **SANCTIONS ALERT**: Found {len(all_matches)} result{'s' if len(all_matches) > 1 else ''} in OFAC database\n"
        result_text += f"*(Minimum match threshold: 83%)*\n\n"
        
        # Display all matches
        for i, match in enumerate(all_matches):
            result_text += f"**Match #{i + 1}: {match['name']}**\n"
            result_text += f"• **Match Score**: {match['score']}\n"
            if match['address']:
                result_text += f"• **Address**: {match['address']}\n"
            result_text += f"• **Entity Type**: {match['entity_type']}\n"
            result_text += f"• **Programs**: {match['programs']}\n"
            result_text += f"• **List**: {match['list_type']}\n"
            result_text += f"\n"
        
        # Add summary and recommendation
        result_text += "---\n\n"
        
        # Check for 100% matches specifically
        perfect_matches = [match for match in all_matches if match['score_num'] == 100]
        
        if perfect_matches:
            # Special message for 100% matches
            result_text += f"**Summary**: Subject returned a 100% match - escalate to management immediately.\n\n"
            result_text += "**Recommendation**: 🚨 **ESCALATE TO MANAGEMENT IMMEDIATELY** - Perfect match detected in OFAC sanctions database."
        elif high_confidence_matches > 0:
            result_text += f"**Summary**: {high_confidence_matches} match{'es' if high_confidence_matches > 1 else ''} (83%+) detected. Entity has strong similarity to sanctioned individuals/entities.\n\n"
            result_text += "**Recommendation**: ⛔ **DO NOT PROCEED** - Conduct thorough manual review and legal consultation before any business relationship."
        else:
            result_text += f"**Summary**: No matches found above 83% similarity threshold.\n\n"
            result_text += "**Recommendation**: ✅ **PROCEED** - No significant OFAC sanctions concerns identified."
        
        # Add direct search URL
        encoded_query = urllib.parse.quote(query)
        direct_url = f"https://sanctionssearch.ofac.treas.gov/?search={encoded_query}"
        result_text += f"\n\n**Direct OFAC Search URL**: {direct_url}"
        
        # Track OFAC search time
        ofac_time = time.time() - start_time
        if 'performance_metrics' in st.session_state:
            st.session_state.performance_metrics['api_response_times'].append(ofac_time)
        logging.info(f"OFAC search completed in {ofac_time:.2f}s")
        
        return result_text
        
    except requests.exceptions.Timeout:
        ofac_time = time.time() - start_time
        logging.warning(f"OFAC search timed out after {ofac_time:.2f}s")
        return "❌ OFAC search timed out. Please try again."
    except requests.exceptions.RequestException as e:
        ofac_time = time.time() - start_time
        logging.error(f"OFAC connection error after {ofac_time:.2f}s: {str(e)}")
        return f"❌ Error connecting to OFAC database: {str(e)}"
    except Exception as e:
        ofac_time = time.time() - start_time
        logging.error(f"OFAC search error after {ofac_time:.2f}s: {str(e)}")
        return f"❌ Error during OFAC search: {str(e)}"

# PDF generation function removed - no longer needed

def search_with_comprehensive(company_name, model=DEFAULT_COMPREHENSIVE_MODEL):
    """Comprehensive search combining Perplexity AI research with OFAC sanctions screening"""
    client_ip = get_client_ip()
    logging.info(f"SEARCH_START - IP: {client_ip} - Comprehensive search for '{company_name}' using model '{model}'")
    try:
        # First, perform OFAC sanctions search
        logging.info(f"Starting comprehensive search for {company_name} - OFAC phase")
        ofac_result = search_with_ofac(company_name)
        
        # Analyze OFAC results for high-risk matches
        ofac_summary = ""
        high_risk_sanctions = False
        
        if "🚨" in ofac_result or "SANCTIONS ALERT" in ofac_result:
            high_risk_sanctions = True
            # Extract key information from OFAC results
            if "Found" in ofac_result:
                count_match = re.search(r'Found (\d+) matches', ofac_result)
                if count_match:
                    match_count = count_match.group(1)
                    ofac_summary = f"CRITICAL: {match_count} matches found in OFAC sanctions databases. "
                else:
                    ofac_summary = "CRITICAL: Multiple matches found in OFAC sanctions databases. "
            else:
                ofac_summary = "CRITICAL: Entity appears on OFAC sanctions lists. "
            
            # Check for 100% matches first, then other high-confidence matches
            if "100%" in ofac_result:
                ofac_summary += "PERFECT MATCH DETECTED (100% similarity). ESCALATE TO MANAGEMENT IMMEDIATELY."
            elif "Score: 9" in ofac_result or "Score: 8" in ofac_result:
                ofac_summary += "High-confidence matches detected (83%+ similarity). RED FLAG ALERT."
        else:
            ofac_summary = "No matches found in OFAC sanctions databases."
        
        # Now perform Perplexity search with OFAC context
        model_config = PERPLEXITY_MODELS.get(model, {})
        logging.info(f"Starting comprehensive search for {company_name} - Perplexity phase using model '{model}' ({model_config.get('name', 'Unknown Model')})")
        
        # Build domain instruction using preferred domains for comprehensive search
        domain_instruction = f"\n\nIMPORTANT: Prioritize information from these authoritative domains first: {', '.join(PREFERRED_DOMAINS)}. Search these regulatory, financial, and news sources before other sources as they provide the most reliable AML-relevant information."
        
        # Enhanced prompt that includes OFAC context
        enhanced_prompt = f"""
You are an expert AML (Anti-Money Laundering) analyst conducting comprehensive due diligence research on "{company_name}".{domain_instruction}

OFAC SANCTIONS SCREENING RESULTS:
{ofac_summary}

Based on the OFAC screening results above and your research, provide a comprehensive AML assessment with the following structure:

## Subject Summary
- Basic subject information and business activities
- Key executives and ownership structure
- Geographic presence and operations

### Subject Websites
IMPORTANT: You MUST include this section. List all official websites for the subject as clickable URLs (e.g., https://example.com). If no websites are found, state: 'This search did not identify potential websites for the subject.'

## AML Risk Assessment

### OFAC Sanctions Analysis
- Incorporate the OFAC screening results above
- If any OFAC matches with 100% similarity were found, this requires IMMEDIATE ESCALATION TO MANAGEMENT
- If any OFAC matches with 83%+ similarity scores were found, this is a CRITICAL RED FLAG
- All matches above 83% similarity threshold require thorough verification
- Explain the implications of any sanctions matches

### Negative News & Compliance Issues
- Money laundering allegations or investigations
- Regulatory violations and enforcement actions
- Criminal investigations or prosecutions
- Suspicious transaction reports or regulatory scrutiny
- Politically Exposed Persons (PEP) connections
- High-risk jurisdiction operations

### Financial Crime Indicators
- Unusual transaction patterns
- Shell company characteristics
- Complex ownership structures
- Offshore entity connections
- Cash-intensive business models

### Regulatory History
- Banking license issues
- Compliance violations
- Regulatory sanctions or penalties
- Supervisory actions

## Summary & Overall Assessment
Based on your analysis, provide a clear summary of key findings regarding this entity:
- Key risk factors identified (if any)
- Notable compliance or regulatory issues
- Current status of any ongoing matters
- Overall risk profile based on available information
- OFAC sanctions screening results and implications

IMPORTANT: If OFAC matches were found, clearly state the match details, similarity scores, and implications. Present findings objectively without making business recommendations.

Provide specific examples and cite your sources. Be thorough but concise.
"""

        # Call Perplexity with enhanced prompt using optimized parameters
        # For Sonar models, system prompts are ignored - combine everything in user message
        if "sonar" in model.lower():
            full_prompt = f"As an expert AML analyst, provide thorough, factual analysis with objective findings summary and proper citations.\n\n{enhanced_prompt}"
            messages = [
                {"role": "user", "content": full_prompt}
            ]
        else:
            # For other models, use system + user structure
            messages = [
                {"role": "system", "content": "You are an expert AML analyst. Provide thorough, factual analysis with objective findings summary and proper citations."},
                {"role": "user", "content": enhanced_prompt}
            ]
        
        # Build API parameters based on model type
        max_tokens = PERPLEXITY_MODELS.get(model, {}).get("max_tokens", 2000)
        api_params = {
            "model": model,
            "messages": messages,
            "max_tokens": max_tokens
        }
        
        # Note: reasoning_effort parameter not yet supported in current OpenAI client version
        # Will be added when the client is updated to support this parameter
        if "deep-research" in model.lower():
            logging.info(f"Deep research model selected for comprehensive search - enhanced analysis will be performed")
        
        # Only add temperature for non-Sonar models (real-time search models don't use temperature)
        if "sonar" not in model.lower():
            api_params["temperature"] = 0.1
        
        logging.info(f"Calling Perplexity API with model: {model} (max_tokens: {max_tokens})...")
        response = openai_client.chat.completions.create(**api_params)
        logging.info(f"Perplexity API call completed using model: {model}")
        
        answer = response.choices[0].message.content
        
        # Extract citations from the response
        citations = extract_citations(response.choices[0].message if response.choices else None, response)
        
        # Add OFAC as a citation with direct search URL
        citations.append({
            'title': 'OFAC Sanctions List Search',
            'url': 'https://sanctionssearch.ofac.treas.gov/'
        })
        
        # Add direct search URL for specific query
        encoded_query = urllib.parse.quote(company_name)
        search_url = f"https://sanctionssearch.ofac.treas.gov/?search={encoded_query}"
        citations.append({
            'title': f'Direct OFAC Search Results for "{company_name}"',
            'url': search_url
        })
        
        # No longer extracting recommendations - providing objective findings only
        recommendation = None
        
        # Append OFAC details to the answer
        full_answer = answer + "\n\n## OFAC Sanctions Screening Details\n\n" + ofac_result
        
        return {
            "status": "success",
            "error": None,
            "answer": full_answer,
            "citations": citations,
            "recommendation": recommendation
        }
        
    except Exception as e:
        logging.error(f"Error in comprehensive search for {company_name}: {str(e)}", exc_info=True)
        return {
            "status": "failed",
            "error": f"Comprehensive search failed: {str(e)}",
            "answer": None,
            "citations": [],
            "recommendation": None
        }

# Initialize OpenAI client for content analysis
@st.cache_resource
def get_openai_client():
    """Get or create OpenAI client for content analysis"""
    openai_api_key, source = load_api_key('OPENAI_API_KEY')
    if not openai_api_key:
        logging.warning(f"OpenAI API key not found. Source: {source}")
        return None
    
    try:
        http_client = get_http_client()
        client = OpenAI(api_key=openai_api_key, http_client=http_client)
        logging.info(f"OpenAI client initialized successfully from {source}")
        return client
    except Exception as e:
        logging.error(f"Failed to initialize OpenAI client: {str(e)}")
        return None

def analyze_content_findings(answer):
    """Use GPT-4.1-mini to determine if there are negative findings in the content"""
    if not answer or not answer.strip():
        return False  # No content means no negative findings
    
    # Check if response is too short to be meaningful
    if len(answer.strip()) < 20:
        return False  # Too short to contain meaningful negative findings
    
    try:
        openai_analysis_client = get_openai_client()
        if not openai_analysis_client:
            logging.error("OpenAI client not available for content analysis")
            return False  # Conservative fallback
        
        analysis_prompt = f"""Analyze this AML research content and determine if there are any negative findings about the entity.

Content:
{answer}

Look for any negative information such as:
- Sanctions or OFAC listings
- Criminal investigations or charges
- Regulatory violations or fines
- Money laundering allegations
- Fraud or financial crimes
- Legal proceedings or lawsuits
- Compliance violations
- Suspicious activities

Respond with exactly one word:
- "NEGATIVE" if any negative findings exist
- "CLEAN" if no negative findings are present

Response:"""

        response = openai_analysis_client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "user", "content": analysis_prompt}
            ],
            max_tokens=5,
            temperature=0
        )
        
        ai_result = response.choices[0].message.content.strip().upper()
        logging.info(f"GPT-4.1-mini analysis result for content analysis: '{ai_result}' (Original: '{response.choices[0].message.content}')")
        return ai_result == "NEGATIVE"
        
    except Exception as e:
        logging.error(f"Error in GPT-4.1-nano content analysis: {str(e)}")
        # Conservative fallback - assume no negative findings if analysis fails
        return False

def add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a paragraph"""
    try:
        # Add the hyperlink relationship
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        
        # Escape XML special characters in the text
        import html
        escaped_text = html.escape(text)
        
        # Create hyperlink XML
        hyperlink_xml = f'''
        <w:hyperlink r:id="{r_id}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
            <w:r>
                <w:rPr>
                    <w:color w:val="0563C1"/>
                    <w:u w:val="single"/>
                    <w:sz w:val="22"/>
                </w:rPr>
                <w:t>{escaped_text}</w:t>
            </w:r>
        </w:hyperlink>
        '''
        
        hyperlink_element = parse_xml(hyperlink_xml)
        paragraph._element.append(hyperlink_element)
        
    except Exception as e:
        logging.warning(f"Failed to create clickable hyperlink for {url}: {e}")
        # Fallback to styled text that looks like a link
        hyperlink_run = paragraph.add_run(text)
        hyperlink_run.font.size = Pt(11)  # Regular text size
        hyperlink_run.font.color.rgb = RGBColor(5, 99, 193)  # Word's default link color
        hyperlink_run.underline = True

def add_text_with_urls(paragraph, text):
    """Add text to paragraph, converting URLs to clickable hyperlinks"""
    import re
    
    # URL pattern to match http/https URLs
    url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+[^\s<>"{}|\\^`\[\].,;!?]'
    
    # Find all URLs in the text
    urls = list(re.finditer(url_pattern, text))
    
    if not urls:
        # No URLs found, add as regular text
        paragraph.add_run(text)
        return
    
    # Process text with URLs
    last_end = 0
    
    for match in urls:
        # Add text before the URL
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        
        # Add the URL as a hyperlink
        url = match.group()
        add_hyperlink(paragraph, url, url)
        
        last_end = match.end()
    
    # Add any remaining text after the last URL
    if last_end < len(text):
        paragraph.add_run(text[last_end:])

def generate_word_document(results_data):
    """Generate a single Word document containing all subjects with summary table and detailed sections"""
    
    # Create new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('AML Due Diligence Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add generation date
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Empty line
    
    # Pre-analyze all content to determine if there are negative findings
    for result in results_data:
        if result['status'] == 'success':
            company_name = result.get('name', 'Unknown')
            answer = result.get('answer', '')
            logging.info(f"Starting content analysis for company: {company_name}")
            has_negative_findings = analyze_content_findings(answer)
            result['has_negative_findings'] = has_negative_findings
            logging.info(f"Content analysis result for {company_name}: {'NEGATIVE' if has_negative_findings else 'CLEAN'}")
    
    # Create summary table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Set column widths
    table.columns[0].width = Inches(2.5)  # Name
    table.columns[1].width = Inches(2.0)  # CLEAR  
    table.columns[2].width = Inches(2.0)  # Internet Search
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Name'
    header_cells[1].text = 'CLEAR'
    header_cells[2].text = 'Internet Search'
    
    # Format header row with light blue background
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add light blue background color to header cells
        shading_elm = parse_xml(r'<w:shd {} w:fill="D9E2F3"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Add data rows
    for result in results_data:
        if result['status'] == 'success':
            row_cells = table.add_row().cells
            row_cells[0].text = result['name']
            
            # Determine CLEAR status - consistent across all recommendations
            recommendation = result.get('recommendation', 'N/A')
            clear_status = 'N/A'
            
            row_cells[1].text = clear_status
            
            # Determine Internet Search status based on GPT-4.1-nano analysis
            has_negative_findings = result.get('has_negative_findings', False)
            
            # Simple logic: if negative findings exist, show "See below", otherwise "No negative news"
            if has_negative_findings:
                internet_search = 'See below'
            else:
                internet_search = 'No negative news'
                
            row_cells[2].text = internet_search
            
            # Center align all cells
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Empty line after table
    
    # Add detailed sections for each subject
    for i, result in enumerate(results_data):
        if result['status'] == 'success':
            # Add subject heading as H1 - large and noticeable
            heading = doc.add_heading(f'{result["name"]} Due Diligence Summary:', level=1)
            
            # Extract and format the answer content
            answer = result.get('answer', '')
            recommendation = result.get('recommendation', 'N/A')
            
            # Add the complete answer content
            if answer and answer.strip():
                # Get the negative findings analysis result
                has_negative_findings = result.get('has_negative_findings', False)
                
                # Process content line by line to handle different header levels
                lines = answer.split('\n')
                current_paragraph = []
                
                for line in lines:
                    line_stripped = line.strip()
                    
                    # Debug: Log header detection
                    if line_stripped.startswith('#'):
                        logging.info(f"Header detected: '{line_stripped}' for {result['name']}")
                    
                    # Handle different markdown header levels (check longest patterns first)
                    if line_stripped.startswith('#### '):
                        # Process any accumulated paragraph content
                        if current_paragraph:
                            para_text = '\n'.join(current_paragraph).strip()
                            if para_text:
                                # Remove markdown formatting but preserve URLs
                                para_text = para_text.replace('**', '').replace('*', '')
                                
                                # Create paragraph and add content with URL detection
                                para = doc.add_paragraph()
                                add_text_with_urls(para, para_text)
                            current_paragraph = []
                        
                        # Add H4 heading
                        header_text = line_stripped[5:].strip()
                        logging.info(f"Creating H4 heading: '{header_text}'")
                        doc.add_heading(header_text, level=4)
                        
                    elif line_stripped.startswith('### '):
                        # Process any accumulated paragraph content
                        if current_paragraph:
                            para_text = '\n'.join(current_paragraph).strip()
                            if para_text:
                                # Remove markdown formatting but preserve URLs
                                para_text = para_text.replace('**', '').replace('*', '')
                                
                                # Create paragraph and add content with URL detection
                                para = doc.add_paragraph()
                                add_text_with_urls(para, para_text)
                            current_paragraph = []
                        
                        # Add H3 heading
                        header_text = line_stripped[4:].strip()
                        logging.info(f"Creating H3 heading: '{header_text}'")
                        doc.add_heading(header_text, level=3)
                        
                    elif line_stripped.startswith('## '):
                        # Process any accumulated paragraph content
                        if current_paragraph:
                            para_text = '\n'.join(current_paragraph).strip()
                            if para_text:
                                # Remove markdown formatting but preserve URLs
                                para_text = para_text.replace('**', '').replace('*', '')
                                
                                # Create paragraph and add content with URL detection
                                para = doc.add_paragraph()
                                add_text_with_urls(para, para_text)
                            current_paragraph = []
                        
                        # Add H2 heading
                        header_text = line_stripped[3:].strip()
                        logging.info(f"Creating H2 heading: '{header_text}'")
                        doc.add_heading(header_text, level=2)
                        
                    elif line_stripped.startswith('# '):
                        # Process any accumulated paragraph content
                        if current_paragraph:
                            para_text = '\n'.join(current_paragraph).strip()
                            if para_text:
                                # Remove markdown formatting but preserve URLs
                                para_text = para_text.replace('**', '').replace('*', '')
                                
                                # Create paragraph and add content with URL detection
                                para = doc.add_paragraph()
                                add_text_with_urls(para, para_text)
                            current_paragraph = []
                        
                        # Add H3 heading for single # (make it bold and larger)
                        header_text = line_stripped[2:].strip()
                        heading_para = doc.add_paragraph()
                        heading_run = heading_para.add_run(header_text)
                        heading_run.bold = True
                        heading_run.font.size = Pt(14)  # Larger than normal text
                        
                    elif line_stripped == '':
                        # Empty line - if we have accumulated content, add paragraph and start new one
                        if current_paragraph:
                            para_text = '\n'.join(current_paragraph).strip()
                            if para_text:
                                # Remove markdown formatting but preserve URLs
                                para_text = para_text.replace('**', '').replace('*', '')
                                
                                # Create paragraph and add content with URL detection
                                para = doc.add_paragraph()
                                add_text_with_urls(para, para_text)
                            current_paragraph = []
                            
                    else:
                        # Regular content line
                        current_paragraph.append(line)
                
                # Process any remaining paragraph content
                if current_paragraph:
                    para_text = '\n'.join(current_paragraph).strip()
                    if para_text:
                        # Remove markdown formatting but preserve URLs
                        para_text = para_text.replace('**', '').replace('*', '')
                        
                        # Create paragraph and add content with URL detection
                        para = doc.add_paragraph()
                        add_text_with_urls(para, para_text)
                    
            else:
                # No answer content at all
                doc.add_paragraph("This search yielded no general background information on the subject.")
                doc.add_paragraph("This search yielded no findings that associate the subject with financial crimes or negative news.")
            
            doc.add_paragraph()  # Empty line
            
            # Add sources section with compact formatting
            sources_para = doc.add_paragraph()
            sources_run = sources_para.add_run('[Sources]')
            sources_run.bold = True
            sources_run.font.size = Pt(10)  # Smaller header
            
            # Add citations if available
            citations = result.get('citations', [])
            if citations:
                for j, citation in enumerate(citations, 1):
                    title = citation.get('title', 'Unknown Source')
                    url = citation.get('url', '#')
                    
                    # Create compact citation paragraph with smaller font
                    citation_para = doc.add_paragraph()
                    citation_para.space_after = Pt(3)  # Reduce spacing after paragraph
                    
                    # Add citation number
                    citation_num_run = citation_para.add_run(f"{j}. ")
                    citation_num_run.font.size = Pt(9)
                    
                    # Make the title itself a clickable hyperlink
                    if url != '#' and url:
                        add_hyperlink(citation_para, url, title)
                    else:
                        # If no URL, just add the title as plain text
                        title_run = citation_para.add_run(title)
                        title_run.font.size = Pt(9)
            else:
                no_sources_para = doc.add_paragraph("No specific sources cited.")
                no_sources_run = no_sources_para.runs[0]
                no_sources_run.font.size = Pt(9)
            
            # Add spacing between subjects
            if i < len(results_data) - 1:
                doc.add_paragraph()
                doc.add_paragraph()
    
    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# --- Streamlit App UI ---
# Professional header with Axos logo
st.markdown("""
<div style="text-align: center; padding: 2rem 0;">
    <img src="data:image/jpeg;base64,{logo_base64}" style="max-height: 80px; margin-bottom: 1rem;" alt="Axos Bank">
    <h1 style="color: #0066CC; font-size: 2.5rem; margin-bottom: 0.5rem; margin-top: 1rem;">
        AML Research Platform
    </h1>
    <div style="margin-top: 1rem;">
        <span style="background-color: #e8f4fd; color: #0066CC; padding: 0.5rem 1rem; border-radius: 20px; font-weight: 600;">
            Version {version}
        </span>
    </div>
</div>
""".format(version=APP_VERSION, logo_base64=get_axos_logo_base64()), unsafe_allow_html=True)



# Create a clean two-column layout for settings
col1, col2 = st.columns([1, 1])

with col1:
    # Search engine selection
    search_engine = st.selectbox(
        "🔍 **Search Method**",
        ["Comprehensive (AI + OFAC)", "AI Research Only", "OFAC Sanctions Only"],
        index=0,
        help="Choose your research approach"
    )

with col2:
    # Model selection (only show if AI is involved)
    if "AI" in search_engine or "Comprehensive" in search_engine:
        model_keys = list(PERPLEXITY_MODELS.keys())
        default_index = model_keys.index(DEFAULT_PERPLEXITY_MODEL) if DEFAULT_PERPLEXITY_MODEL in model_keys else 0
        selected_model = st.selectbox(
            "🤖 **AI Model**",
            model_keys,
            format_func=lambda x: PERPLEXITY_MODELS[x]["name"],
            index=default_index,
            help="Choose AI research depth"
        )
        
        # Show cost warning for deep research model
        if "deep-research" in selected_model.lower():
            st.caption("ℹ️ 💰💰💰 Deep Research model provides exhaustive analysis but has higher token costs")
        
        reasoning_effort = "medium"  # Default (reasoning_effort not yet supported in client)
    else:
        selected_model = DEFAULT_OFAC_FALLBACK_MODEL  # Default for OFAC-only
        reasoning_effort = "medium"

# Check if we need AI and if client is available
needs_ai = "AI" in search_engine or "Comprehensive" in search_engine
if needs_ai and not openai_client:
    st.error("⚠️ Perplexity API client required but not initialized. Please check your API key configuration.")
    st.stop()

st.markdown("---")

# Main input section with enhanced styling
st.markdown("""
<div style="background-color: white; padding: 2rem; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.08); margin-bottom: 1rem;">
    <h3 style="color: #0066CC; margin-top: 0;">📋 Subject Research Queue</h3>
</div>
""", unsafe_allow_html=True)

company_names_input = st.text_area(
    "**Subject Names** (one per line)",
    height=150,
    placeholder="Example:\nGazprom\nAxos Financial\nMicrosoft\nJohn Smith\nShell Company Ltd",
    key="company_names",
    help="Enter company names, individual names, or entity names for AML screening"
)

# Show subject count
if company_names_input:
    lines = [line.strip() for line in company_names_input.split('\n') if line.strip()]
    line_count = len(lines)
    if line_count == 1:
        st.caption(f"📝 {line_count} subject entered")
    else:
        st.caption(f"📝 {line_count} subjects entered")
else:
    line_count = 0
    st.info("👆 Enter subject names above to begin")

# Word document output (only format available)
st.markdown("### 📄 **Output Format**")
st.info("📄 Word document will be generated for download")
output_format = "Single Word Document"  # Fixed format


st.markdown("---")

# Generate Reports Button
if line_count > 0:
    if st.button("🚀 **Generate AML Reports**", type="primary"):
        overall_start_time = time.time()  # Track overall processing time
        subject_names = [name.strip() for name in company_names_input.split('\n') if name.strip()]
        
        # Show processing info
        model_name = PERPLEXITY_MODELS.get(selected_model, {}).get("name", selected_model) if "AI" in search_engine or "Comprehensive" in search_engine else "N/A"
        
        # Show generation mode
        st.info(f"🔄 Processing {len(subject_names)} subjects using **{search_engine}** {f'with **{model_name}**' if model_name != 'N/A' else ''}")
        
        # Update statistics
        st.session_state.search_stats['total_subjects'] += len(subject_names)
        
        st.session_state.results_list = []
        progress_bar = st.progress(0)
        status_text = st.empty()
        total_names = len(subject_names)

        for i, name in enumerate(subject_names):
            search_start_time = time.time()  # Track individual search time
            status = "failed"
            error_message = "Processing not started."
            recommendation = None
            answer = ""
            citations = []
            api_time = 0
            
            # Update status
            status_text.text(f"🔍 Analyzing {name} ({i+1}/{total_names})...")
            
            with st.spinner(f"Processing {name}..."):
                # Log the request before processing
                log_user_request(search_engine, name, selected_model)
                
                # Perform search based on selected engine - SAME FOR BOTH FORMATS
                if search_engine == "Comprehensive (AI + OFAC)":
                    result = search_with_comprehensive(name, selected_model)
                    if result["status"] == "success":
                        status = "success"
                        error_message = None
                        recommendation = result.get("recommendation")
                        answer = result.get("answer", "")
                        citations = result.get("citations", [])
                    else:
                        status = "failed"
                        error_message = result["error"]
                        recommendation = None
                        answer = ""
                        citations = []
                
                elif search_engine == "AI Research Only":
                    result = search_with_perplexity(name, selected_model)
                    if result["status"] == "success":
                        status = "success"
                        error_message = None
                        recommendation = result.get("recommendation")
                        answer = result.get("answer", "")
                        citations = result.get("citations", [])
                    else:
                        status = "failed"
                        error_message = result["error"]
                        recommendation = None
                        answer = ""
                        citations = []
                
                elif search_engine == "OFAC Sanctions Only":
                    ofac_result = search_with_ofac(name)
                    if "❌" in ofac_result:
                        status = "failed"
                        error_message = ofac_result
                        recommendation = None
                        answer = ""
                        citations = []
                    else:
                        status = "success"
                        error_message = None
                        answer = ofac_result
                        citations = [{'title': 'OFAC Sanctions Database', 'url': 'https://sanctionssearch.ofac.treas.gov/'}]
                        recommendation = None  # No longer providing recommendations
                
                # PDF generation removed - only Word documents supported
            
            # Track individual search time
            search_time = time.time() - search_start_time
            st.session_state.performance_metrics['search_times'].append(search_time)
            
            # Log the completion status
            if status == "success":
                logging.info(f"SEARCH_SUCCESS - IP: {get_client_ip()} - '{name}' completed successfully in {search_time:.2f}s with recommendation: {recommendation}")
                st.session_state.search_stats['successful_searches'] += 1
                
                # Check for risk alerts
                if answer and ("OFAC" in answer and ("match" in answer.lower() or "100%" in answer or "83%" in answer)):
                    st.session_state.search_stats['ofac_hits'] += 1
                    st.session_state.search_stats['risk_alerts'] += 1
                elif answer and analyze_content_findings(answer):
                    st.session_state.search_stats['risk_alerts'] += 1
            else:
                logging.error(f"SEARCH_FAILED - IP: {get_client_ip()} - '{name}' failed after {search_time:.2f}s: {error_message}")
                st.session_state.search_stats['failed_searches'] += 1
            
            # Add search to history
            st.session_state.performance_metrics['search_history'].append({
                'id': f"{datetime.now().timestamp()}_{i}",
                'name': name,
                'time': search_time,
                'status': status,
                'timestamp': datetime.now()
            })
            
            st.session_state.results_list.append({
                'name': name,
                'status': status,
                'error_message': error_message,
                'recommendation': recommendation,
                'answer': answer,
                'citations': citations,
                'search_time': search_time
            })
            
            progress_bar.progress((i + 1) / total_names)

        # Track overall processing time
        overall_time = time.time() - overall_start_time
        st.session_state.performance_metrics['total_processing_time'] += overall_time
        st.session_state.performance_metrics['last_search_time'] = overall_time
        st.session_state.performance_metrics['total_searches'] += len(subject_names)
        
        st.success(f"✅ **Processing Complete!** Total time: {overall_time:.1f}s ({overall_time/len(subject_names):.1f}s per subject)")
        progress_bar.empty()
        status_text.empty()
else:
    st.button("🚀 **Generate AML Reports**", disabled=True, help="Enter subject names first")

# --- Display Results Status ---
if st.session_state.results_list:
    st.markdown("---")
    st.markdown("### 📊 **Processing Results**")
    
    # Performance breakdown
    with st.expander("⏱️ **Performance Breakdown**", expanded=True):
        perf_col1, perf_col2, perf_col3 = st.columns(3)
        
        with perf_col1:
            if st.session_state.performance_metrics['search_times']:
                avg_search = sum(st.session_state.performance_metrics['search_times']) / len(st.session_state.performance_metrics['search_times'])
                st.metric("Average Search Time", f"{avg_search:.2f}s")
            else:
                st.metric("Average Search Time", "N/A")
        
        with perf_col2:
            if st.session_state.performance_metrics['api_response_times']:
                avg_api = sum(st.session_state.performance_metrics['api_response_times']) / len(st.session_state.performance_metrics['api_response_times'])
                st.metric("Average API Response", f"{avg_api:.2f}s")
            else:
                st.metric("Average API Response", "N/A")
        
        with perf_col3:
            doc_time = st.session_state.performance_metrics.get('document_generation_time', 0)
            if doc_time > 0:
                st.metric("Document Generation", f"{doc_time:.2f}s")
            else:
                st.metric("Document Generation", "N/A")
        
        # Show individual search times
        if st.session_state.results_list:
            st.markdown("**Individual Search Times:**")
            for result in st.session_state.results_list:
                search_time = result.get('search_time', 0)
                if result['status'] == 'success':
                    st.success(f"✅ {result['name']}: {search_time:.2f}s")
                else:
                    st.error(f"❌ {result['name']}: {search_time:.2f}s (failed)")
    
    # Word document mode - show summary and generate single document
    successful_results = [res for res in st.session_state.results_list if res['status'] == 'success']
    failed_results = [res for res in st.session_state.results_list if res['status'] != 'success']
    
    if successful_results:
        st.success(f"✅ {len(successful_results)} subjects processed successfully")
    if failed_results:
        st.error(f"❌ {len(failed_results)} subjects failed")
        for result in failed_results:
            st.error(f"• **{result['name']}**: {result.get('error_message', 'Unknown error')}")
    
    # Generate and offer Word document download
    if successful_results:
        st.markdown("---")
        st.markdown("### 📄 **Download Word Document**")
        
        try:
            doc_start_time = time.time()
            logging.info(f"DOCUMENT_GENERATION - IP: {get_client_ip()} - Generating Word document for {len(successful_results)} subjects")
            word_bytes = generate_word_document(successful_results)
            doc_gen_time = time.time() - doc_start_time
            st.session_state.performance_metrics['document_generation_time'] = doc_gen_time
            
            st.download_button(
                label="📥 **Download AML Due Diligence Report (Word)**",
                data=word_bytes,
                file_name=f"AML_Due_Diligence_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )
            
            st.info(f"📋 Document contains summary table and detailed sections for {len(successful_results)} subjects (generated in {doc_gen_time:.1f}s)")
            
        except Exception as e:
            st.error(f"❌ Error generating Word document: {str(e)}")

# Footer
st.markdown("---")
st.markdown(f"**AML Demo v{APP_VERSION}** | Powered by Perplexity AI & OFAC Database (83% match threshold) | Word Document Reports Only") 